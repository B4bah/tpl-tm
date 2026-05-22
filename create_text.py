"""
Скрипт для автоматической генерации пояснительной записки в формате .docx
Требуется установка библиотеки python-docx: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def set_cell_font(cell, font_name='Times New Roman', font_size=10, bold=False):
    """Установка шрифта для ячейки таблицы."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold

def set_cell_text(cell, text, bold=False, font_name='Times New Roman', font_size=10):
    """Запись текста в ячейку с заданным форматированием."""
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold

def add_code_block(doc, code_text):
    """Добавление блока кода с моноширинным шрифтом и серым фоном."""
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles['Normal']
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    # Установка отступа
    paragraph.paragraph_format.left_indent = Cm(1.27)
    for line in code_text.split('\n'):
        run = paragraph.add_run(line + '\n')
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        # Серый фон
        run.font.highlight_color = None  # python-docx не умеет напрямую фон для run; используем shading для параграфа позже
        # альтернатива: задать стиль с заливкой
    # Применим заливку к абзацу
    from docx.oxml.ns import qn
    pPr = paragraph._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F2F2F2',
        qn('w:val'): 'clear'
    })
    pPr.append(shd)

def add_heading_styled(doc, text, level):
    """Добавление заголовка с форматированием."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def create_document():
    doc = Document()

    # Настройка стилей по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Заголовки
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(13)

    # ======= ТИТУЛЬНАЯ ЧАСТЬ =======
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ПОЯСНИТЕЛЬНАЯ ЗАПИСКА')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('к лабораторной работе по дисциплине\n«Теория языков программирования и методы трансляции»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    topic = doc.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic.add_run('Тема: Разработка компилятора учебного языка на фрагмент языка C\nс использованием PLY (Python Lex-Yacc)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()  # пустая строка

    # ======= 1. ЗАДАНИЕ =======
    add_heading_styled(doc, '1. Задание', level=1)

    p = doc.add_paragraph(
        'Разработать компилятор с учебного языка на выходной фрагмент кода на языке C. '
        'Входной язык содержит конструкцию «последовательность операторов цикла с параметром»:'
    )
    add_code_block(doc, '<выражение инициализации> for <условное выражение> <выражение итерации> <оператор>')
    doc.add_paragraph(
        'Тело цикла – оператор присваивания, в левой части которого обязательно присутствует '
        'переменная с индексом (элемент массива). Язык обладает следующими лексико-синтаксическими особенностями:'
    )
    items = [
        'Идентификаторы: длина не более 8 символов, конец определяется разделителем (пробел, скобки и т.д.).',
        'Служебные слова: не выделяются специальными символами, записываются как обычные идентификаторы.',
        'Знаки операций: кодируются служебными словами (например, add вместо +, sub вместо –).',
        'Числа: вещественные десятичные с порядком (экспоненциальная форма), целые константы.',
        'Логические операции и отношения: не равно (ne), дизъюнкция (or), отрицание (not).',
        'Операции над числовыми и символьными величинами: умножение (mul), деление (div), извлечение квадратного корня (sqrt).',
        'Операнды выражений: целые константы, вещественные константы, целые переменные, булевские переменные.',
        'Форма записи выражений: строго префиксная, с произвольным уровнем вложенности.',
        'Индекс массива: произвольное выражение (не только константа или идентификатор).'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'Дополнение согласно п.6 задания: Для обеспечения функциональной полноты языка в состав операций, '
        'помимо указанных, включены операции сравнения (eq, lt, gt, le, ge), конъюнкция (and), '
        'а также арифметические сложение (add) и вычитание (sub). Решение принято по согласованию с '
        'преподавателем для устранения противоречия между составом требуемых операций и необходимостью '
        'построения осмысленных циклов.'
    )

    # ======= 2. ОПИСАНИЕ ЛЕКСЕРА =======
    add_heading_styled(doc, '2. Описание лексера', level=1)

    add_heading_styled(doc, '2.1. Токены', level=2)
    doc.add_paragraph('Перечень лексем, распознаваемых лексером, приведён в таблице 1.')

    # Таблица 1 – Токены
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table1.rows[0].cells
    set_cell_text(hdr[0], 'Токен', bold=True)
    set_cell_text(hdr[1], 'Регулярное выражение / описание', bold=True)
    set_cell_text(hdr[2], 'Значение (если применимо)', bold=True)

    tokens_data = [
        ('ID', '[a-zA-Z_][a-zA-Z_0-9]* (max 8 символов)', 'Строка идентификатора'),
        ('INTEGER_CONST', r'\d+ (не часть вещественного числа)', 'Целое значение (int)'),
        ('REAL_CONST', 'См. алгоритм 2.2 (вещественные с порядком)', 'Вещественное значение (float)'),
        ('ASSIGN', '=', ''),
        ('LBRACKET', '[', ''),
        ('RBRACKET', ']', ''),
        ('FOR', 'for', ''),
        ('AND', 'and', ''),
        ('OR', 'or', ''),
        ('NOT', 'not', ''),
        ('EQ', 'eq', ''),
        ('NE', 'ne', ''),
        ('LT', 'lt', ''),
        ('GT', 'gt', ''),
        ('LE', 'le', ''),
        ('GE', 'ge', ''),
        ('ADD', 'add', ''),
        ('SUB', 'sub', ''),
        ('MUL', 'mul', ''),
        ('DIV', 'div', ''),
        ('SQRT', 'sqrt', ''),
        ('TRUE', 'true', 'Булево значение True'),
        ('FALSE', 'false', 'Булево значение False')
    ]
    for t in tokens_data:
        row = table1.add_row().cells
        set_cell_text(row[0], t[0])
        set_cell_text(row[1], t[1])
        set_cell_text(row[2], t[2])

    # Подпись таблицы
    caption1 = doc.add_paragraph('Таблица 1 – Токены входного языка')
    caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption1.runs[0].italic = True if caption1.runs else False
    if not caption1.runs:
        run = caption1.add_run('Таблица 1 – Токены входного языка')
        run.italic = True
    else:
        caption1.runs[0].italic = True

    add_heading_styled(doc, '2.2. Алгоритмы распознавания токенов', level=2)

    steps = [
        'Пробельные символы: пробел и табуляция игнорируются; символ перевода строки увеличивает внутренний счётчик строк для диагностики.',
        'Идентификаторы: распознаются по регулярному выражению [a-zA-Z_][a-zA-Z_0-9]*. После выделения лексема проверяется на вхождение в таблицу служебных слов – если совпадение есть, возвращается соответствующий токен. Иначе проводится контроль длины: при превышении 8 символов генерируется ошибка, токен пропускается.',
        'Вещественные константы: используются регулярные выражения для трёх форм:',
        '   - число.дробная_часть (с необязательным порядком): \\d+\\.\\d*([eE][+-]?\\d+)?',
        '   - число без дробной части, но с порядком: \\d+[eE][+-]?\\d+',
        '   - число с ведущей точкой: \\.\\d+([eE][+-]?\\d+)?',
        'Эти выражения объединены в одно правило. Распознанная лексема преобразуется в тип float.',
        'Целые константы: последовательность цифр, которая не была поглощена правилом вещественной константы. Значение – int.',
        'Булевы константы: слова true и false преобразуются в булевы значения True / False соответственно.',
        'Неизвестные символы: вызывают лексическую ошибку с указанием строки, проблемный символ пропускается для продолжения анализа.'
    ]
    for step in steps:
        doc.add_paragraph(step)

    doc.add_paragraph('Лексер построен с помощью библиотеки PLY (ply.lex).')

    # ======= 3. ОПИСАНИЕ ПАРСЕРА =======
    add_heading_styled(doc, '3. Описание парсера', level=1)

    add_heading_styled(doc, '3.1. Грамматика языка', level=2)
    doc.add_paragraph('Грамматика, реализованная в парсере, приведена в форме Бэкуса–Наура (БНФ). '
                      'Нетерминалы выделены полужирным шрифтом, терминалы – моноширинным.')

    grammar = (
        "<program>       : <for_list>\n\n"
        "<for_list>      : <for_stmt>\n"
        "                | <for_stmt> <for_list>\n\n"
        "<for_stmt>      : <expression> FOR <expression> <expression> <assign_stmt>\n\n"
        "<assign_stmt>   : ID '[' <expression> ']' '=' <expression>\n\n"
        "<expression>    : ADD <expression> <expression>\n"
        "                | SUB <expression> <expression>\n"
        "                | MUL <expression> <expression>\n"
        "                | DIV <expression> <expression>\n"
        "                | AND <expression> <expression>\n"
        "                | OR  <expression> <expression>\n"
        "                | EQ  <expression> <expression>\n"
        "                | NE  <expression> <expression>\n"
        "                | LT  <expression> <expression>\n"
        "                | GT  <expression> <expression>\n"
        "                | LE  <expression> <expression>\n"
        "                | GE  <expression> <expression>\n"
        "                | NOT <expression>\n"
        "                | SQRT <expression>\n"
        "                | ID '=' <expression>\n"
        "                | ID\n"
        "                | INTEGER_CONST\n"
        "                | REAL_CONST\n"
        "                | TRUE\n"
        "                | FALSE\n"
        "                | ID '[' <expression> ']'"
    )
    add_code_block(doc, grammar)

    add_heading_styled(doc, '3.2. Устранение конфликтов', level=2)
    doc.add_paragraph(
        'Использование строго префиксной нотации для всех операций полностью исключает неоднозначность '
        'приоритетов: оператор всегда предшествует операндам, поэтому в грамматике отсутствуют '
        'традиционные правила приоритетов. Вложенность выражений управляется исключительно рекурсивными правилами.'
    )
    doc.add_paragraph(
        'Конструкция assign_stmt (тело цикла) имеет фиксированную структуру '
        'ID \'[\' <expression> \']\' \'=\' <expression>, что позволяет парсеру однозначно отделить её '
        'от обычного выражения. Конфликта с array_ref внутри <expression> не возникает, так как '
        'в контексте тела цикла парсер ожидает именно assign_stmt.'
    )
    doc.add_paragraph(
        'Парсер, построенный на основе LALR-грамматики библиотекой PLY (ply.yacc), не генерирует '
        'предупреждений о конфликтах сдвига/свёртки.'
    )

    add_heading_styled(doc, '3.3. Диагностика и локализация ошибок', level=2)
    doc.add_paragraph(
        'Лексические ошибки (недопустимый символ, превышение допустимой длины идентификатора) '
        'фиксируются с выводом сообщения, содержащего проблемный символ/лексему и номер строки. '
        'Ошибочный токен пропускается, анализ продолжается.'
    )
    doc.add_paragraph(
        'Синтаксические ошибки перехватываются функцией p_error, определённой в парсере. '
        'Выводится сообщение с указанием токена, на котором возникла ошибка, и (при наличии) '
        'номера строки. После ошибки парсер пытается продолжить разбор, однако корректность '
        'последующего AST не гарантируется.'
    )

    # ======= 4. ПРИМЕРЫ ВХОДНЫХ ТЕКСТОВ =======
    add_heading_styled(doc, '4. Примеры входных текстов', level=1)

    add_heading_styled(doc, '4.1. Корректные примеры', level=2)
    doc.add_paragraph(
        'В таблице 2 приведены примеры входных строк и соответствующий им выходной фрагмент на C. '
        'Во всех случаях компиляция проходит успешно.'
    )

    # Таблица 2 – Корректные примеры
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    set_cell_text(hdr2[0], '№', bold=True)
    set_cell_text(hdr2[1], 'Входная строка', bold=True)
    set_cell_text(hdr2[2], 'Выходной C-фрагмент', bold=True)

    examples_good = [
        ('1',
         'i = 0 for lt i 10 i = add i 1 a[i] = mul i 2',
         'for (i = 0; (i < 10); i = (i + 1)) {\n    a[i] = (i * 2);\n}'),
        ('2',
         'n = 5 for ne n 0 n = sub n 1 arr[ sub n 1 ] = div 10 n',
         'for (n = 5; (n != 0); n = (n - 1)) {\n    arr[(n - 1)] = (10 / n);\n}'),
        ('3',
         'x = 1.5e2 for and gt x 100 lt x 200 x = add x 0.5 data[ sqrt x ] = not false',
         'for (x = 150.0; ((x > 100) && (x < 200)); x = (x + 0.5)) {\n    data[sqrt(x)] = !(0);\n}'),
        ('4',
         'k = 0 for lt mul k 3 30 k = add k 1 b[k] = add mul 2 k div k 2',
         'for (k = 0; ((k * 3) < 30); k = (k + 1)) {\n    b[k] = ((2 * k) + (k / 2));\n}'),
        ('5',
         'a = 0 for lt a 3 a = add a 1 x[a] = a b = 10 for gt b 0 b = sub b 1 y[b] = b',
         'for (a = 0; (a < 3); a = (a + 1)) {\n    x[a] = a;\n}\nfor (b = 10; (b > 0); b = (b - 1)) {\n    y[b] = b;\n}')
    ]

    for ex in examples_good:
        row = table2.add_row().cells
        set_cell_text(row[0], ex[0])
        set_cell_text(row[1], ex[1])
        # Для кода используем Courier New
        cell = row[2]
        cell.text = ''
        p = cell.paragraphs[0]
        for line in ex[2].split('\n'):
            run = p.add_run(line + '\n')
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    caption2 = doc.add_paragraph('Таблица 2 – Корректные входные данные и результат трансляции')
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption2.runs[0].italic = True

    add_heading_styled(doc, '4.2. Ошибочные примеры', level=2)
    doc.add_paragraph(
        'В таблице 3 приведены входные строки, содержащие ошибки, и реакция компилятора.'
    )

    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr3 = table3.rows[0].cells
    set_cell_text(hdr3[0], '№', bold=True)
    set_cell_text(hdr3[1], 'Входная строка', bold=True)
    set_cell_text(hdr3[2], 'Ожидаемое поведение / сообщения', bold=True)

    examples_bad = [
        ('1', 'longident = 1 for lt i 5 i = add i 1 arr[i] = 0',
         'Lexer error: identifier \'longident\' too long (max 8 chars)\nДалее – синтаксическая ошибка, в C-коде инициализация искажена.'),
        ('2', 'i = 0 for i 10 i = add i 1 a[i] = 0',
         'Syntax error at token \'=\' (line 1)\nПропущена операция сравнения – парсер не может построить цикл.'),
        ('3', 'i = 0 for lt i 10 a[i] = i',
         'Syntax error at token \'=\' (line 1)\nОтсутствует выражение итерации.'),
        ('4', 'i = 0 for lt i 10 i = add i 1 arr[i] =',
         'Syntax error at end of input\nНезавершённое присваивание.'),
        ('5', 'x = 1.2.3 for lt x 5 x = add x 1 a[0] = x',
         'Syntax error at token \'=\' (line 1)\nЛексер не распознаёт число с двумя точками как одну лексему.'),
        ('6', 'i = 0 for less i 10 i = add i 1 a[i] = i',
         'less распознаётся как ID, а не оператор; ошибка в структуре for.'),
        ('7', 'for = 0 for lt for 10 for = add for 1 arr[for] = for',
         'Ключевое слово for нельзя использовать как идентификатор; возникает ошибка структуры.')
    ]
    for ex in examples_bad:
        row = table3.add_row().cells
        set_cell_text(row[0], ex[0])
        set_cell_text(row[1], ex[1])
        set_cell_text(row[2], ex[2])

    caption3 = doc.add_paragraph('Таблица 3 – Ошибочные входные данные и диагностические сообщения')
    caption3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption3.runs[0].italic = True

    # ======= 5. ЗАКЛЮЧЕНИЕ =======
    add_heading_styled(doc, '5. Заключение', level=1)
    doc.add_paragraph(
        'Разработанный компилятор полностью удовлетворяет заданным требованиям: реализованы '
        'лексический, синтаксический анализ и генерация кода на C. Поддерживаются все предписанные '
        'типы операций, префиксная форма записи, неограниченная вложенность выражений, '
        'индекс-выражение, вещественные константы с порядком и ограничение длины идентификаторов. '
        'Обеспечена диагностика лексических и синтаксических ошибок.'
    )

    # Сохранение
    doc.save('Пояснительная_записка.docx')
    print('Документ "Пояснительная_записка.docx" успешно создан.')

if __name__ == '__main__':
    create_document()